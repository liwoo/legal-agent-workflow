#!/usr/bin/env python3
"""Generate the Contract Triage end-to-end testing how-to guide as a .docx.

Run from the e2e/ folder (with the agent venv active):
    python tools/build_guide.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
E2E = HERE.parent
SHOTS = E2E / "artifacts" / "screenshots"
OUT = E2E / "Contract-Triage-E2E-Testing-Guide.docx"

ACCENT = RGBColor(0xE1, 0x55, 0x1E)   # Northgate orange
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x66, 0x66, 0x66)
CODE_BG = "F3F4F6"

outcome = {}
op = SHOTS.parent / "outcome.json"
if op.exists():
    outcome = json.loads(op.read_text())

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK


def _shade(cell_or_para, color_hex):
    el = cell_or_para._tc if hasattr(cell_or_para, "_tc") else cell_or_para._p
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    (el.find(qn("w:tcPr")) if hasattr(cell_or_para, "_tc") else _para_pr(cell_or_para)).append(shd)


def _para_pr(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    return pPr


def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = INK
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def code(text):
    p = doc.add_paragraph()
    _shade(p, CODE_BG)
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(text.strip("\n").split("\n")):
        r = p.add_run(("\n" if i else "") + line)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return p


def shot(name, caption, width=6.3):
    path = SHOTS / name
    if not path.exists():
        body(f"[missing screenshot: {name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    cap.paragraph_format.space_after = Pt(10)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 2"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ── Cover ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
tr = title.add_run("Contract Triage")
tr.bold = True
tr.font.size = Pt(28)
tr.font.color.rgb = ACCENT
title.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
sr = sub.add_run("End-to-End Testing Guide — for Testers")
sr.font.size = Pt(15)
sr.font.color.rgb = INK
sub.paragraph_format.space_after = Pt(2)

meta = doc.add_paragraph()
mr = meta.add_run(
    "A Playwright walkthrough that creates a contract in the console, watches the "
    "agent triage it live, confirms the outcome, and observes the run in Langfuse.\n"
    "Version 1.0 · Northgate Systems Ltd (internal)"
)
mr.italic = True
mr.font.size = Pt(10)
mr.font.color.rgb = MUTED

doc.add_paragraph()

# ── 1. What this proves ──────────────────────────────────────────────────────
h1("1.  What this test proves")
body(
    "This end-to-end (E2E) test exercises the whole Contract Triage product in one "
    "pass — the same journey a legal reviewer takes, driven by a real browser and a "
    "real LLM. A single run proves that every layer is wired together correctly:"
)
bullet("opens with the queues already warm (the ten seeded contracts).", "Console — ")
bullet("captures the intake facts and an attached PDF, then submits.", "New Contract form — ")
bullet("the submit button shows the “Triaging…” spinner while the agent runs.", "Loading state — ")
bullet(
    "the Microsoft Agent Framework graph classifies the paper, runs the policy "
    "gates, and reaches a disposition via real gpt-4o calls.",
    "Agent — ",
)
bullet(
    "the new contract appears in the queue it was triaged to (Signed / Review / "
    "Inbox), and its detail view shows the outcome, journey, and reasoning.",
    "Hydration — ",
)
bullet(
    "the run shows up as a span tree — workflow, agent invocation, and each chat "
    "call with prompts, token counts, latency and cost.",
    "Observability — ",
)
body(
    "Because triage is a live model call, the outcome is not fixed: the same NDA may "
    "be signed on one run and escalated on another. The test adapts — it reads the "
    "disposition from the API response and follows the contract to whichever queue it "
    "lands in. A pass means the pipeline works, not that a specific verdict was reached."
)

# ── 2. Prerequisites ─────────────────────────────────────────────────────────
h1("2.  Prerequisites")
table(
    ["Tool", "Version", "Why"],
    [
        ["Docker", "running", "Langfuse trace stack (7 containers)"],
        ["Python + uv", "3.12", "FastAPI backend + agent"],
        ["Node.js", "≥ 20 (22 tested)", "Next.js console + Playwright"],
        ["OpenAI API key", "gpt-4o access", "the agent is LLM-first; no key = no triage"],
    ],
)
body(
    "The OpenAI key is the one hard requirement. Without it the triage nodes raise "
    "LLMUnavailableError and Langfuse shows no chat spans. Put it in app/agent/.env "
    "(git-ignored):"
)
code("OPENAI_API_KEY=sk-...\nOPENAI_CHAT_MODEL=gpt-4o")

# ── 3. Architecture ──────────────────────────────────────────────────────────
h1("3.  The stack under test")
body(
    "The test drives an already-running stack. These are the moving parts and where "
    "they listen:"
)
table(
    ["Service", "URL", "Role"],
    [
        ["Review console (Next.js)", "http://localhost:3000", "what the test clicks"],
        ["Triage API (FastAPI)", "http://localhost:8000", "persist + run the agent"],
        ["Langfuse", "http://localhost:3001", "trace ingest + UI"],
        ["MinIO (contract PDFs)", "http://localhost:9092", "stores the uploaded PDF"],
        ["MinIO / Postgres / ClickHouse / Redis", "(internal)", "Langfuse backing stores"],
    ],
)
body(
    "The repo ships an Aspire AppHost (host/apphost.mts) that starts all of this with "
    "one command (aspire run). If your Aspire CLI build does not support the TypeScript "
    "AppHost (it may report “Unrecognized app host type”), use the framework-free "
    "harness in the e2e/ folder instead — it reproduces exactly the same containers, "
    "ports and Langfuse keys. The rest of this guide uses that harness.",
)

# ── 4. One-time setup ────────────────────────────────────────────────────────
h1("4.  One-time setup")
numbered("Install the backend (Python) dependencies:", "")
code(
    "cd app/agent\n"
    "uv venv --python 3.12 && source .venv/bin/activate\n"
    "uv pip install -e . --prerelease=allow\n"
    "cd ../.."
)
numbered("Install the frontend dependencies (Node ≥ 20):", "")
code("cd app/frontend && npm install && cd ../..")
numbered("Install Playwright and its browser:", "")
code("cd e2e && npm install && npx playwright install chromium && cd ..")
numbered("Add your OpenAI key to app/agent/.env (see section 2).", "")

# ── 5. Bring up the stack ────────────────────────────────────────────────────
h1("5.  Bring the stack up")
h2("5.1  Start Langfuse + the object stores (Docker)")
code("cd e2e\ndocker compose -f docker-compose.langfuse.yml up -d")
body(
    "Wait ~60–90s on first run while Langfuse migrates its databases and provisions "
    "the Contract Triage project headlessly. It is ready when http://localhost:3001 "
    "returns the sign-in page."
)
h2("5.2  Start the API and the console")
body("From the e2e/ folder, in two terminals:")
code("bash run-api.sh        # FastAPI on :8000, traces → Langfuse\nbash run-frontend.sh   # Next.js console on :3000")
body(
    "run-api.sh sources stack.env, which points the agent's OpenTelemetry exporter at "
    "Langfuse over HTTP (Langfuse does not accept gRPC). On boot the API seeds ten "
    "example contracts and eagerly triages them, so the queues are warm — this is the "
    "first batch of traces you will see in Langfuse. Confirm all three are up:"
)
code(
    'curl -s localhost:8000/api/health          # {\"status\":\"ok\"}\n'
    "curl -s -o /dev/null -w \"%{http_code}\\n\" localhost:3000\n"
    "curl -s -o /dev/null -w \"%{http_code}\\n\" localhost:3001"
)

# ── 6. Run the test ──────────────────────────────────────────────────────────
h1("6.  Run the end-to-end test")
code("cd e2e\nnpm test")
body(
    "A headless Chromium runs the walkthrough below and writes a screenshot for each "
    "stage to e2e/artifacts/screenshots/. Expected console output:"
)
code(
    "Running 1 test using 1 worker\n"
    "➡  Triaged CR-2026-0XX → end_state=signed_desk_edits → queue=signed (Signed — desk edits)\n"
    "  ✓  1 …  create a contract, triage it, and see where it lands\n"
    "  1 passed"
)

h2("Stage 1 — the console loads with warm queues")
shot("01-console-loaded.png", "The Inbox on load. Counts (Inbox / Review / Signed) come from the ten seeded contracts.")

h2("Stage 2 — fill the New Contract form and attach the PDF")
body(
    "The test clicks New Contract, fills the intake facts (counterparty, what arrived, "
    "the sender's ask, sector) and attaches e2e/fixtures/sample-contract.pdf — a mutual "
    "NDA from “Meridian Freight Solutions Ltd”."
)
shot("02-form-filled.png", "The intake form, completed, with the sample PDF attached.")

h2("Stage 3 — submit; the “Triaging…” loading state")
body(
    "On submit the button flips to a spinner labelled “Triaging…” and stays there while "
    "the agent runs synchronously. This is the loading state the test captures."
)
shot("03-triaging-loading.png", "The form mid-run: “Triaging…” while the agent classifies and gates the contract.")

h2("Stage 4 — the page hydrates into the right queue")
body(
    "When the agent finishes, the dialog closes and the new contract appears in the "
    "queue matching its disposition. In this run it was escalated, landing in Review; "
    "on other runs it signs and lands in Signed."
)
shot("04-queue-hydrated.png", "The triaged contract, newly arrived in its queue with the disposition badge.")

h2("Stage 5 — the detail view: outcome, journey and reasoning")
body(
    "Opening the row shows the full triage result: the journey strip "
    "(Arrived → Reviewed → Checked → Signed/Escalated), the model's plain-English "
    "explanation, the recommended action, and — on the Decision tab — the outcome badge."
)
shot("05-detail-summary.png", "Detail view, Summary tab: the journey and the model's “in a nutshell” reasoning.")
shot("06-detail-decision.png", "Detail view, Decision tab: the outcome and what the assistant recommends.")

# ── 7. Langfuse ──────────────────────────────────────────────────────────────
h1("7.  Observe the run in Langfuse")
body("Every triage run is traced. To see the run the test just created:")
numbered("Open http://localhost:3001 and sign in:", "")
code("email:    admin@northgate.local\npassword: langfuse-admin")
numbered("Open the Contract Triage project → Tracing. Filter Trace Name to triage_contract.", "")
shot("07-langfuse-traces.png", "The Tracing list: one triage_contract (and one workflow.build) trace per run.")
numbered("Open the most recent triage_contract trace to see its span tree.", "")
shot("08-langfuse-trace-detail.png", "The span tree: workflow.run → classify → invoke_agent Triage → chat gpt-4o, with token counts, latency and cost. The metadata pins the trace to the contract id the test created.")
body(
    "The trace's metadata (contract.id, contract.counterparty) matches the CR-2026-0XX "
    "the test created and printed — that is your end-to-end link from a click in the "
    "browser to a costed LLM span in the observability stack."
)

# ── 8. Pass criteria ─────────────────────────────────────────────────────────
h1("8.  What a pass looks like")
bullet("npm test reports “1 passed”.", "")
bullet("Six stage screenshots exist under e2e/artifacts/screenshots/ (01–06).", "")
bullet("e2e/artifacts/outcome.json records the created id, end_state and queue.", "")
bullet("A new triage_contract trace appears in Langfuse with a chat gpt-4o span carrying a token count and cost.", "")
if outcome:
    body(
        f"Reference from the authoring run: contract {outcome.get('contractId')} → "
        f"{outcome.get('endState')} → {outcome.get('landedQueue')} "
        f"({outcome.get('disposition')}).",
        italic=True,
    )

# ── 9. Fixes made ────────────────────────────────────────────────────────────
h1("9.  Two defects found (and fixed) while authoring this test")
body(
    "Running the console against the live API — rather than its offline fixtures — "
    "surfaced two real bugs. Both are fixed in the frontend:"
)
h2("9.1  Live queues rendered empty")
body(
    "The backend labels a contract's queue with its disposition vocabulary "
    "(approved / quarantined / pending), but the console filters on inbox / signed / "
    "review. The initial load trusted the API's value verbatim, so every queue matched "
    "nothing and showed empty — the console only looked populated in offline/fixture "
    "mode. Fix: derive the queue from end_state at the API boundary (src/lib/api.ts), "
    "the same rule create/triage/resolve already used."
)
h2("9.2  Duplicated confidence badge")
body(
    "The detail modal rendered the confidence ScoreBadge twice (a copy-paste "
    "duplicate), so the header read “Medium confidence · Medium confidence”. Fix: "
    "remove the duplicate line in src/components/contract-detail-modal.tsx."
)

# ── 10. Troubleshooting ──────────────────────────────────────────────────────
h1("10.  Troubleshooting")
table(
    ["Symptom", "Cause & fix"],
    [
        ["Queues empty / rows missing", "Frontend fix in §9.1 not applied, or the API returned nothing — check localhost:8000/api/contracts."],
        ["“LLMUnavailableError” in the API log", "OPENAI_API_KEY missing/invalid in app/agent/.env."],
        ["No traces in Langfuse", "API not started via run-api.sh (missing OTEL env), or Langfuse still migrating. Give it 90s and retry."],
        ["ModuleNotFoundError …otlp.proto.grpc", "Set OTEL_EXPORTER_OTLP_PROTOCOL=http/protobuf (already in stack.env) — Langfuse is HTTP-only."],
        ["Test times out waiting for a row", "Triage can take 30–60s; the config allows for it. Confirm the API is reachable from the browser."],
        ["Next.js won't start (engine error)", "Node < 20. Use nvm: nvm use 22. run-frontend.sh does this automatically if a newer Node is installed."],
    ],
)

# ── 11. Harness map ──────────────────────────────────────────────────────────
h1("11.  The e2e/ harness — file map")
table(
    ["Path", "Purpose"],
    [
        ["docker-compose.langfuse.yml", "Langfuse stack + object stores (mirrors host/apphost.mts)"],
        ["stack.env", "OTEL → Langfuse + object-store env for the API"],
        ["run-api.sh / run-frontend.sh", "launch the API and the console against the stack"],
        ["playwright.config.ts", "Playwright config (base URL, timeouts, artifacts)"],
        ["tests/contract-triage.e2e.spec.ts", "the end-to-end walkthrough"],
        ["fixtures/sample-contract.pdf", "the intake PDF the test uploads"],
        ["tools/capture-langfuse.mjs", "logs into Langfuse and screenshots a trace"],
        ["artifacts/screenshots/", "per-stage screenshots (also embedded in this guide)"],
    ],
)

doc.save(str(OUT))
print(f"wrote {OUT}")
